#!/usr/bin/python3

import argparse
from docx import Document
from docx.drawing import Drawing

from docx.text.paragraph import Paragraph

def create_parser():
    parser = argparse.ArgumentParser(
        prog='Formating checker',
    )
    parser.add_argument('filename')
    return parser

def main():
    parser = create_parser()
    args = parser.parse_args()

    filename = args.filename

    #print(filename)    

    document = Document(filename)


    for style in document.styles:
        #print(style.name)
        pass

    for section in document.sections:
        #print(section.header.paragraphs[0].part.element)
        pass

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for s in run.iter_inner_content():
                if(Drawing == type(s)):
                    print(s.part.blob)
                    #print(s)
            #print(type(run.iter_inner_content()))
        #print(paragraph.hyperlinks)
        #print(paragraph.part.blob)
        #print()
        pass

    for table in document.tables:
        #print(table)
        #print(paragraph.part.blob)
        #print()
        pass

    #print(document.part.part.blob)
    # for part in document.part:
    #     print(part.)
    #     #print()
    #     pass

    #print(len(document.sections))

    # doc_elm =document._element

    # toc_paragraph_elms = doc_elm.xpath('.//w:sdt[w:sdtPr[w:docPartObj[w:docPartGallery[@w:val="Table of Contents"]]]]/w:sdtContent/w:p')
    # for toc_paragraph_elm in toc_paragraph_elms:
    #     #print('toc_paragraph_elm',toc_paragraph_elm)
    #     #print('parent',toc_paragraph_elm.getparent())
    #     paragraph = Paragraph(toc_paragraph_elm, toc_paragraph_elm.getparent())
    #     if paragraph:
    #         par_text = paragraph.text
    #         print('par_text',par_text)

if __name__ == "__main__":
    main()